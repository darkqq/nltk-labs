import time
import nltk
from nltk.corpus import wordnet as wn
from nltk.draw import TreeWidget
from nltk.draw.util import CanvasFrame
from tkinter import messagebox, Label, RIGHT, Frame, Tk, LEFT, Text, WORD, END, Menu, Button
from tkinter.filedialog import asksaveasfilename, askopenfilename
from docx import Document

nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')
nltk.download('universal_tagset')
nltk.download('omw-1.4')

root = Tk()
root.option_add('*Dialog.msg.font', 'Helvetica 10')
root.title("Syntax parse tree")
root.resizable(width=False, height=False)

label_frame = Frame(root)
label_frame.pack()

enter_label = Label(label_frame, text='Введите текст', width=15)
enter_label.pack(side=LEFT)

empty_label = Label(label_frame, text='', width=65)
empty_label.pack(side=RIGHT)

enter_text = Text(width=70, height=20, wrap=WORD)
enter_text.pack()

canvas = CanvasFrame(root, width=550, height=180)
canvas.pack()


def docx_parser(docx):
    doc = Document(docx)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text


def open_file_and_input_text():
    file_name = askopenfilename(filetypes=(("Docx files", "*.docx"),))
    if file_name != '':
        text = docx_parser(file_name)
        enter_text.delete(1.0, END)
        enter_text.insert(1.0, text)


def save_docx():
    file = asksaveasfilename(filetypes=(("Docx file", "*.docx"),), defaultextension=("Docx file", "*.docx"))
    doc = Document()

    doc.add_paragraph(enter_text.get(1.0, END))
    doc.save(file)


help_info = """Система семантического анализа естественного языка

Система позволяет провести семантический анализ предложения анлглийского языка, загрузить предложение из файла, а также сохранить предложение в файл в формате docx.

Результат семантического анализ предложения представлен в виде дерева, узлами которого являются лексемы и их определение, синонимы, антонимы, гипонимы и гиперонимы.

Для проведения семантического анализа необходимо ввести текст в верхнее поле и затем нажать кнопку "Создать".
Для сохранение необходимо нажать кнопку "Сохранить", в появившемся окне выбрать нужный файл или задать имя новому файлу.
Для открытия словаря необходимо нажать пункт меню "Файл", в появившемся окне выбрать нужный файл.
"""


def information():
    messagebox.askquestion("Help", help_info, type='ok')


def draw_semantic_tree():
    canvas.canvas().delete("all")
    start = time.time()
    text = enter_text.get(1.0, END)
    text = text.replace('\n', '')
    if text != '':
        sentences = nltk.sent_tokenize(text)
        enter_text.insert(END, '\nPlease waiting. Semantic tree is drawing')
        root.update()
        result = '(S '
        for sent in sentences:
            sent = sent.replace('.', '')
            sent = sent.replace(',', '')
            sent = sent.replace('?', '')
            doc = nltk.word_tokenize(sent)
            result_sent = '(SENT '
            for word in doc:
                result_sent += get_word_semantic(word)
            result_sent += ')'
            result += result_sent
        result += ')'
        result = nltk.tree.Tree.fromstring(result)
        widget = TreeWidget(canvas.canvas(), result)
        canvas.add_widget(widget, 50, 10)

    finish = time.time()
    delta = finish - start
    enter_text.delete(enter_text.search('\nPlease waiting. Semantic tree is drawing', 1.0, END), END)
    print('draw tree: ', delta)


def get_word_semantic(word: str) -> str:
    start = time.time()
    if len(wn.synsets(word)) == 0:
        return '(WS (W ' + word + '))'
    result = '(WS (W ' + word + ') (DEF ' + wn.synsets(word)[0].definition().replace(' ', '_') + ')'
    synonyms, antonyms, hyponyms, hypernyms = [], [], [], []
    word = wn.synsets(word)
    syn_app = synonyms.append
    ant_app = antonyms.append
    he_app = hyponyms.append
    hy_app = hypernyms.append
    for synset in word:
        for lemma in synset.lemmas():
            syn_app(lemma.name())
            if lemma.antonyms():
                ant_app(lemma.antonyms()[0].name())
    for hyponym in word[0].hyponyms():
        he_app(hyponym.name())
    for hypernym in word[0].hypernyms():
        hy_app(hypernym.name())
    if len(synonyms):
        result += ' (SYN '
        for synonym in synonyms:
            result += synonym + ' '
    if len(antonyms):
        result += ') (ANT '
        for antonym in antonyms:
            result += antonym + ' '
    if len(hyponyms):
        result += ') (HY '
        for hyponym in hyponyms:
            result += hyponym + ' '
    if len(hypernyms):
        result += ') (HE '
        for hypernym in hypernyms:
            result += hypernym + ' '
    result += '))'
    print('get word semantic', time.time() - start)
    return result


main_menu = Menu(root)
main_menu.add_command(label='Файл', command=open_file_and_input_text)
main_menu.add_command(label='Помощь', command=information)
root.config(menu=main_menu)

button1 = Button(text="Создать", command=draw_semantic_tree)
button1.pack(side=LEFT)

button2 = Button(text="Сохранить", command=save_docx)
button2.pack(side=RIGHT)

root.mainloop()
