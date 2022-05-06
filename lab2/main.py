import time
import nltk
from nltk.draw import TreeWidget
from nltk.draw.util import CanvasFrame
from tkinter import messagebox, Label, RIGHT, Frame, Tk, LEFT, Text, WORD, END, Menu, Button
from tkinter.filedialog import asksaveasfilename, askopenfilename
from docx import Document

nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')
nltk.download('universal_tagset')

root = Tk()
root.option_add('*Dialog.msg.font', 'Helvetica 10')
root.title("Syntax parse tree")
root.resizable(width=False, height=False)

label_frame = Frame(root)
label_frame.pack()

enter_label = Label(label_frame, text='Введите текст', width=10)
enter_label.pack(side=LEFT)

empty_label = Label(label_frame, text='', width=65)
empty_label.pack(side=RIGHT)

enter_text = Text(width=70, height=20, wrap=WORD)
enter_text.pack()

canvas = CanvasFrame(root, width=1000, height=500)
canvas.pack()


def docx_parser(docx):
    doc = Document(docx)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text


def open_file_and_input_text():
    file_name = askopenfilename(filetypes=(("Docx files", "*.docx"),))
    if file_name != '':
        text = docx_parser(file_name)
        enter_text.delete(1.0, END)
        enter_text.insert(1.0, text)


def save_docx():
    file = asksaveasfilename(filetypes=(("Docx file", "*.docx"),), defaultextension=("Docx file", "*.docx"))
    doc = Document()

    doc.add_paragraph(enter_text.get(1.0, END))
    doc.save(file)



grammar = r"""
        P: {<PRT|ADP>}
        V: {<VERB>}
        N: {<NOUN|PRON>}
        NP: {<N|NP|P>+<ADJ|NUM|DET>+}
        NP: {<ADJ|NUM|DET>+<N|NP|P>+}
        PP: {<P><NP>|<NP><P>}
        VP: {<NP|N><V>}
        VP: {<VP><NP|N||ADV>}
        VP: {<NP|N|ADV><VP>}
        VP: {<VP><PP|P>}
        """


def draw_syntax_tree():
    canvas.canvas().delete("all")
    start = time.time()
    text = enter_text.get(1.0, END)
    text = text.replace('\n', '')
    text = text.replace(',', '')
    text = text.replace('.', '')
    if text != '':
        doc = nltk.word_tokenize(text)
        doc = nltk.pos_tag(doc, tagset='universal')
        cp = nltk.RegexpParser(grammar)
        result = cp.parse(doc)
        widget = TreeWidget(canvas.canvas(), result)
        canvas.add_widget(widget, 50, 10)

    finish = time.time()
    delta = finish - start
    print(delta)


main_menu = Menu(root)
main_menu.add_command(label='Файл', command=open_file_and_input_text)

root.config(menu=main_menu)

button1 = Button(text="Создать", command=draw_syntax_tree)
button1.pack(side=LEFT)

button2 = Button(text="Сохранить", command=save_docx)
button2.pack(side=RIGHT)

root.mainloop()
